#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build jajiga-data-report.docx — a human-readable, DATA-ONLY Persian report.
Contains every dataset of the jajiga-tracker project: no analysis, no algorithms.

Sources: jajiga_complete_dataset.json (lossless aggregation of all data files).
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent

# Dataset location: 1) CLI arg  -> 2) env JAJIGA_DATASET  -> 3) H:\hermes outputs  -> 4) project root
import os
import sys

_CANDIDATES = [
    Path(sys.argv[1]) if len(sys.argv) > 1 else None,
    Path(os.environ["JAJIGA_DATASET"]) if os.environ.get("JAJIGA_DATASET") else None,
    Path(r"H:\hermes outputs\jajiga_complete_dataset.json"),
    BASE / "jajiga_complete_dataset.json",
]
AGG = next((p for p in _CANDIDATES if p is not None and p.exists()), _CANDIDATES[-1])
OUT = BASE / "jajiga-data-report.docx"

FONT = "B Nazanin"
HEADER_FILL = "D9E2F3"
ALT_FILL = "F2F2F2"

doc = Document()

# ---------------------------------------------------------------- page setup
for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
rpr = normal.element.get_or_add_rPr()
rf = rpr.get_or_add_rFonts()
for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rf.set(qn(attr), FONT)

# ---------------------------------------------------------------- helpers
def rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    return p


def para(text="", bold=False, size=10.5, align=None, before=0, after=4, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return rtl(p)


def heading(text, level=1):
    sizes = {0: 22, 1: 15, 2: 12.5, 3: 11.5}
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)
    p.paragraph_format.space_before = Pt(16 if level <= 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return rtl(p)


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def fmt(v):
    if v is None:
        return ""
    if isinstance(v, bool):
        return "بله" if v else "خیر"
    if isinstance(v, float) and v == int(v):
        v = int(v)
    if isinstance(v, (int, float)):
        return f"{v:,}"
    return str(v)


def fa_gender(v):
    return {"male": "مرد", "female": "زن"}.get(v, v or "")


def yes_no(v):
    if v in (1, True, "1", "true"):
        return "دارد"
    if v in (0, False, "0", "false"):
        return "ندارد"
    return fmt(v)


def make_table(headers, rows, widths=None, font_size=8.5, alt=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        rtl(p)
        shade(c, HEADER_FILL)
        if widths:
            c.width = Cm(widths[i])
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(fmt(val))
            r.font.size = Pt(font_size)
            rtl(p)
            if widths:
                c.width = Cm(widths[i])
        if alt and ridx % 2 == 1:
            for c in cells:
                shade(c, ALT_FILL)
    return t


def kv_table(pairs, font_size=9, widths=(7, 10)):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for k, v in pairs:
        cells = t.add_row().cells
        for ci, (txt, bold) in enumerate(((k, True), (fmt(v), False))):
            c = cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(txt)
            r.bold = bold
            r.font.size = Pt(font_size)
            rtl(p)
            c.width = Cm(widths[ci])
    return t


def compact_dict(d):
    if not d:
        return ""
    return "؛ ".join(f"{k}: {v}" if v not in (None, "") else str(k) for k, v in d.items())


def compact_list(l):
    if not l:
        return ""
    return "، ".join(str(x) for x in l)


# ---------------------------------------------------------------- load data
agg = json.load(open(AGG, encoding="utf-8"))
hosts = agg["hosts_db"]["hosts"]
all_cabins = agg["all_cabins"]
supply = agg["supply"]
room_dates = agg["room_dates"]
pricing_json = agg["pricing"]["json"]
pricing_csv = agg["pricing"]["csv"]
raw_pricing = agg["raw_pricing"]
snapshots = agg["snapshots"]

# join supply rooms (village/est_date) with hosts_db rooms (rich fields)
supply_by_id = {str(r["id"]): r for r in supply["rooms"]}
rich_rooms = []
for h in hosts:
    for r in h.get("rooms", []):
        s = supply_by_id.get(str(r["id"]), {})
        rich_rooms.append(
            {
                **r,
                "host_id": h.get("id"),
                "host_name": h.get("name"),
                "village": s.get("village", ""),
                "est_date": s.get("est_date", ""),
                "j_est": s.get("j_est", ""),
                "status": s.get("status", ""),
            }
        )

# ---------------------------------------------------------------- title page
para("", after=60)
para("گزارش جامع داده‌های پروژه جاجیگا", bold=True, size=26,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
para("بابل‌کنار — بازار اقامتگاه‌های بوم‌گردی و کلبه", bold=True, size=15,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
para("مستندسازی کامل دیتاست — فقط داده، بدون تحلیل", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
para("منبع: jajiga_complete_dataset.json", size=10.5,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=40)

# ---------------------------------------------------------------- TOC
heading("فهرست مطالب", 1)
toc_items = [
    "۱. معرفی دیتاست",
    "۲. خلاصه کلی (شمارنده‌ها)",
    "۳. میزبان‌ها — ۳۴۶ میزبان",
    "۴. اقامتگاه‌ها — ۴۶۷ اقامتگاه (از پروفایل میزبان‌ها)",
    "۵. روند عرضه — ۶۸ ماه (تایم‌لاین)",
    "۶. کلبه‌های ۶ روستا — ۱۵۷ کلبه (با وضعیت اشغال)",
    "۷. دیتاست قیمت‌گذاری — ۱۰۸ کلبه / ۴ روستا",
    "۸. تاریخ‌های تخمینی اقامتگاه‌ها — ۴۶۷ رکورد",
    "۹. جزئیات کامل API — ۱۱۸ اقامتگاه (خام)",
    "پیوست ۱. توضیحات میزبان‌ها",
    "پیوست ۲. اسنپ‌شات‌های تاریخی عرضه",
]
for it in toc_items:
    para(it, size=11.5, after=3)

doc.add_page_break()

# ---------------------------------------------------------------- 1. intro
heading("۱. معرفی دیتاست", 1)
para("این سند شامل تمام داده‌های جمع‌آوری‌شده در پروژه jajiga-tracker است: بازار اقامتگاه‌های "
     "بوم‌گردی و کلبه در منطقه بابل‌کنار (مازندران). داده‌ها از وب‌سایت جاجیگا و API آن بدون نیاز به "
     "احراز هویت جمع‌آوری شده‌اند و در این گزارش به‌صورت خوانا برای انسان ارائه می‌شوند.")
para("محدوده جغرافیایی: بابل‌کنار، بابل — مازندران. نوع اقامتگاه: کلبه و اقامتگاه بوم‌گردی (cottage).")
para("هر فصل از این گزارش یک منبع داده مستقل را پوشش می‌دهد و همه فیلدهای آن منبع بدون حذف "
     "شامل شده است. تصاویر در این گزارش ثبت نمی‌شوند؛ اما نتیجه تحلیل آن‌ها (تاریخ تخمینی اقامتگاه‌ها) "
     "در فصل‌های ۴ و ۸ و ۹ آمده است.", after=8)

# ---------------------------------------------------------------- 2. summary
heading("۲. خلاصه کلی (شمارنده‌ها)", 1)
counts = agg["meta"]["counts"]
summary_rows = [
    ("تعداد میزبان‌ها (پروفایل کامل)", counts["hosts_db_hosts"]),
    ("تعداد اقامتگاه‌ها (از پروفایل میزبان‌ها)", counts["supply_rooms"]),
    ("کلبه‌های ۶ روستای منتخب (all-cabins)", counts["all_cabins_total"]),
    ("روستاهای منتخب", "، ".join(counts["all_cabins_villages"])),
    ("ماه‌های پوشش‌داده‌شده در روند عرضه", counts["supply_months"]),
    ("اقامتگاه‌های دارای تاریخ تخمینی", counts["room_dates_rooms"]),
    ("اسنپ‌شات‌های تاریخی", counts["snapshots"]),
    ("کلبه‌های دیتاست قیمت‌گذاری (JSON)", counts["pricing_cabins_json"]),
    ("کلبه‌های دیتاست قیمت‌گذاری (CSV)", counts["pricing_cabins_csv"]),
    ("فایل‌های خام API (raw_pricing)", counts["raw_pricing_total"]),
]
kv_table(summary_rows, font_size=10)

para("", after=4)
heading("تفکیک فایل‌های خام API", 2)
raw_rows = [[v, c] for v, c in counts["raw_pricing_files"].items()]
make_table(["دایرکتوری (روستا)", "تعداد فایل"], raw_rows, widths=(6, 6), font_size=10)

doc.add_page_break()

# ---------------------------------------------------------------- 3. hosts
heading("۳. میزبان‌ها — ۳۴۶ میزبان", 1)
para("منبع: data/hosts-babolkenar.json — پروفایل کامل میزبان‌های بابل‌کنار (از API کاربر جاجیگا).")

host_rows = []
for h in hosts:
    host_rows.append(
        [
            h.get("id"),
            h.get("name"),
            fa_gender(h.get("gender")),
            h.get("verified"),
            h.get("member_since"),
            h.get("rooms_count"),
            h.get("active_rooms_count"),
            h.get("total_success_books"),
            h.get("host_level"),
            h.get("price_range"),
            h.get("avg_price"),
            h.get("accept_rate"),
            h.get("response_time_min"),
            h.get("communication_rate"),
        ]
    )
make_table(
    ["شناسه", "نام", "جنسیت", "تأیید", "عضویت", "تعداد اقامتگاه", "فعال", "رزرو موفق",
     "سطح", "محدوده قیمت (تومان)", "میانگین قیمت", "نرخ پذیرش", "زمان پاسخ (دقیقه)", "امتیاز ارتباط"],
    host_rows,
    widths=(1.4, 2.0, 0.9, 1.0, 1.9, 1.3, 1.0, 1.2, 1.2, 2.0, 1.5, 1.2, 1.5, 1.4),
    font_size=8,
)

para("", after=2)
para("توضیحات میزبان‌ها (۱۵۱ میزبان دارای توضیح) در پیوست ۱ آمده است.", size=9.5)

doc.add_page_break()

# ---------------------------------------------------------------- 4. rooms
heading("۴. اقامتگاه‌ها — ۴۶۷ اقامتگاه (از پروفایل میزبان‌ها)", 1)
para("منبع: فیلد rooms در data/hosts-babolkenar.json + روستا و تاریخ تخمینی از data/supply-data.json.")
room_rows = []
for r in rich_rooms:
    room_rows.append(
        [
            r.get("id"),
            r.get("title"),
            r.get("village"),
            r.get("host_name"),
            r.get("price"),
            r.get("success_books"),
            r.get("bedrooms"),
            r.get("floor_area"),
            r.get("guests"),
            r.get("class"),
            r.get("rating"),
            r.get("reviews"),
            r.get("is_plus"),
            r.get("is_instant"),
            r.get("is_clean"),
            r.get("discount"),
            r.get("est_date"),
        ]
    )
make_table(
    ["شناسه", "عنوان", "روستا", "میزبان", "قیمت (تومان)", "رزرو", "اتاق‌خواب", "متراژ",
     "ظرفیت", "کلاس", "امتیاز", "نظرات", "ممتاز", "رزرو فوری", "نظافت", "تخفیف٪", "تخمین عضویت"],
    room_rows,
    widths=(1.3, 3.2, 1.4, 1.7, 1.5, 1.0, 0.9, 0.9, 0.8, 1.2, 0.8, 0.8, 0.8, 0.9, 0.8, 0.8, 1.5),
    font_size=7.5,
)

doc.add_page_break()

# ---------------------------------------------------------------- 5. timeline
heading("۵. روند عرضه — ۶۸ ماه (تایم‌لاین)", 1)
para("منبع: data/supply-data.json → months — تعداد میزبان‌ها و اقامتگاه‌های جدید در هر ماه و تجمعی آن‌ها.")
month_rows = []
for m in supply["months"]:
    month_rows.append(
        [m.get("key"), m.get("j"), m.get("hosts_new"), m.get("rooms_new"),
         m.get("hosts_cum"), m.get("rooms_cum")]
    )
make_table(
    ["ماه (میلادی)", "ماه (شمسی)", "میزبان جدید", "اقامتگاه جدید", "میزبان تجمعی", "اقامتگاه تجمعی"],
    month_rows,
    widths=(2.6, 2.6, 2.6, 2.6, 2.6, 2.6),
    font_size=9,
)

doc.add_page_break()

# ---------------------------------------------------------------- 6. cabins
heading("۶. کلبه‌های ۶ روستا — ۱۵۷ کلبه (با وضعیت اشغال)", 1)
para("منبع: data/all-cabins.json → villages — کلبه‌های منتخب در ۶ روستا با وضعیت اشغال ۳۰ روز گذشته.")
village_rows = []
for vname, cabins in all_cabins["villages"].items():
    for c in cabins:
        village_rows.append(
            [
                vname,
                c.get("id"),
                c.get("title"),
                c.get("price"),
                c.get("rooms"),
                c.get("floor"),
                c.get("guests"),
                c.get("rating"),
                c.get("reviews"),
                c.get("success_books"),
                c.get("occupancy_30"),
                c.get("occupancy_30_unavailable"),
                c.get("occupancy_30_total"),
                c.get("pool"),
                c.get("jacuzzi"),
                c.get("host"),
                c.get("active"),
                c.get("own"),
            ]
        )
make_table(
    ["روستا", "شناسه", "عنوان", "قیمت (تومان)", "اتاق", "طبقه", "ظرفیت", "امتیاز", "نظرات",
     "رزرو موفق", "اشغال ۳۰", "ناموجود ۳۰", "مجموع ۳۰", "استخر", "جکوزی", "میزبان", "فعال", "کلبه خود"],
    village_rows,
    widths=(1.5, 1.2, 3.0, 1.5, 0.8, 0.8, 0.9, 0.8, 0.8, 1.1, 1.0, 1.1, 1.0, 0.8, 0.8, 1.4, 0.7, 0.9),
    font_size=7.5,
)

doc.add_page_break()

# ---------------------------------------------------------------- 7. pricing
heading("۷. دیتاست قیمت‌گذاری — ۱۰۸ کلبه / ۴ روستا", 1)
para("منبع: data/pricing/pricing-dataset.json — داده‌های کامل قیمت‌گذاری فاز ۱ "
     "(جمع‌آوری داده؛ الگوریتم و تحلیل در این گزارش وجود ندارد).")
para("این فصل فقط داده را ثبت می‌کند: فاکتورهای هر کلبه، امکانات، برچسب‌ها، تخفیف‌ها، نظرات و اطلاعات میزبان.")

pricing_rows = []
for c in pricing_json:
    pricing_rows.append(
        [
            c.get("id"),
            c.get("title"),
            c.get("village"),
            c.get("bedrooms"),
            c.get("floor_area"),
            c.get("guest_number"),
            c.get("min_price"),
            c.get("extra_price"),
            c.get("success_books"),
            c.get("rating"),
            c.get("reviews"),
            c.get("occupancy_30"),
            c.get("pool"),
            c.get("jacuzzi"),
            "، ".join(c.get("properties") or []),
            "، ".join(c.get("features") or []),
        ]
    )
make_table(
    ["شناسه", "عنوان", "روستا", "اتاق", "متراژ", "ظرفیت", "قیمت پایه", "قیمت اضافه",
     "رزرو", "امتیاز", "نظرات", "اشغال ۳۰", "استخر", "جکوزی", "برچسب‌ها", "امکانات"],
    pricing_rows,
    widths=(1.2, 2.6, 1.3, 0.8, 0.9, 0.9, 1.3, 1.2, 0.9, 0.8, 0.8, 0.9, 0.8, 0.8, 2.2, 3.0),
    font_size=7.5,
)

para("", after=4)
heading("جزئیات کامل هر کلبه", 2)
para("در ادامه، همه فیلدهای باقی‌مانده برای هر یک از ۱۰۸ کلبه ثبت شده است.", size=9.5)

for c in pricing_json:
    heading(f"کلبه {c.get('id')} — {c.get('title', '')}", 3)
    pairs = [
        ("شناسه", c.get("id")),
        ("عنوان", c.get("title")),
        ("آدرس صفحه", c.get("url")),
        ("وضعیت", c.get("status")),
        ("روستا", c.get("village")),
        ("اسلاگ روستا", c.get("village_slug")),
        ("تعداد اتاق‌خواب", c.get("bedrooms")),
        ("زیربنا (متر)", c.get("floor_area")),
        ("زمین (متر)", c.get("land_area")),
        ("تعداد طبقات", c.get("floors_count")),
        ("ظرفیت عادی (نفر)", c.get("guest_number")),
        ("حداکثر ظرفیت (نفر)", c.get("max_guest_number")),
        ("چیدمان خواب", compact_list(c.get("sleep_arrange"))),
        ("نوع اقامتگاه", compact_list(c.get("types"))),
        ("منطقه", compact_list(c.get("regions"))),
        ("حداقل اقامت (شب)", c.get("stays_min")),
        ("حداکثر اقامت (شب)", c.get("stays_max")),
        ("قیمت پایه هر شب (تومان)", c.get("min_price")),
        ("قیمت هر نفر اضافه (تومان)", c.get("extra_price")),
        ("سیاست لغو", c.get("cancellation_policy")),
        ("اقامتگاه ممتاز (پلاس)", c.get("is_plus")),
        ("رزرو فوری", c.get("is_instant")),
        ("نشان نظافت", c.get("is_clean")),
        ("اقامتگاه جدید", c.get("is_new")),
        ("برچسب‌ها (ویژگی‌ها)", compact_list(c.get("properties"))),
        ("امکانات (کلیدها)", compact_list(c.get("features"))),
        ("تعداد امکانات", c.get("features_count")),
        ("شرح امکانات", compact_dict(c.get("feature_desc"))),
        ("رزرو موفق", c.get("success_books")),
        ("امتیاز کلی", c.get("rating")),
        ("تعداد نظرات", c.get("reviews")),
        ("امتیاز دقت", c.get("rating_accuracy")),
        ("امتیاز ارتباط", c.get("rating_communication")),
        ("امتیاز نظافت", c.get("rating_cleanliness")),
        ("امتیاز موقعیت", c.get("rating_location")),
        ("امتیاز ورود", c.get("rating_checkin")),
        ("امتیاز ارزش", c.get("rating_value")),
        ("اخیراً فعال", c.get("is_recently")),
        ("میزبان نامطلوب", c.get("is_bad_host")),
        ("تخفیف فعلی (٪)", c.get("current_discount_percent")),
        ("تخفیف فعلی (جزئیات)", c.get("current_discount")),
        ("تخفیف‌های مدت اقامت", compact_list(c.get("discounts"))),
        ("تعداد عکس", c.get("pictures_count")),
        ("عکس ۳۶۰", c.get("vr_photo")),
        ("ویدیو", c.get("video_url")),
        ("عرض جغرافیایی", c.get("geo", {}).get("lat") if isinstance(c.get("geo"), dict) else None),
        ("طول جغرافیایی", c.get("geo", {}).get("lng") if isinstance(c.get("geo"), dict) else None),
        ("شناسه میزبان", c.get("host_id")),
        ("نام میزبان", c.get("host_name")),
        ("جنسیت میزبان", fa_gender(c.get("host_gender"))),
        ("تاریخ عضویت میزبان", c.get("host_created_at")),
        ("نرخ پذیرش میزبان", c.get("host_accept_rate")),
        ("زمان پاسخ میزبان (دقیقه)", c.get("host_response_time")),
        ("امتیاز ارتباط میزبان", c.get("host_communication_rate")),
        ("اشغال ۳۰ روز", c.get("occupancy_30")),
        ("ناموجود ۳۰ روز", c.get("occupancy_30_unavailable")),
        ("مجموع ۳۰ روز", c.get("occupancy_30_total")),
        ("استخر", yes_no(c.get("pool"))),
        ("جکوزی", yes_no(c.get("jacuzzi"))),
        ("کلبه خود کاربر", c.get("own")),
    ]
    kv_table(pairs, font_size=8.5)
    para("", after=2)

doc.add_page_break()

# ---------------------------------------------------------------- 8. room dates
heading("۸. تاریخ‌های تخمینی اقامتگاه‌ها — ۴۶۷ رکورد", 1)
para("منبع: data/supply/room-dates.json — تاریخ تخمینی اقامتگاه از اولین عکس و اولین نظر محاسبه شده "
     "(بازه: تاریخ عضویت میزبان تا اولین عکس/نظر، هرکدام زودتر).")
date_rows = []
for rid, rd in room_dates.items():
    date_rows.append(
        [
            rd.get("id"),
            rd.get("title"),
            rd.get("host_name"),
            rd.get("host_created_at"),
            rd.get("first_photo"),
            rd.get("photo_count"),
            rd.get("first_review"),
            rd.get("est_date"),
            rd.get("status"),
        ]
    )
make_table(
    ["شناسه", "عنوان", "میزبان", "عضویت میزبان", "اولین عکس", "تعداد عکس", "اولین نظر",
     "تاریخ تخمینی", "وضعیت"],
    date_rows,
    widths=(1.3, 3.4, 2.0, 2.0, 1.8, 1.2, 1.8, 1.8, 1.2),
    font_size=8,
)

doc.add_page_break()

# ---------------------------------------------------------------- 9. raw api details
heading("۹. جزئیات کامل API — ۱۱۸ اقامتگاه (خام)", 1)
para("منبع: داده‌های خام GET /api/room/{id} از ۵ دایرکتوری — کامل‌ترین منبع داده برای هر اقامتگاه. "
     "تصاویر ثبت نشده‌اند؛ اما تعداد عکس، تاریخ‌ها و توضیحات عکس‌ها آمده است. "
     "تاریخ تخمینی از فایل room-dates به هر اقامتگاه متصل شده است.")

for village_key, items in raw_pricing.items():
    if not items:
        continue
    village_names = {
        "seydkola": "سیدکلا",
        "gonehkola": "گونه کلا",
        "quran_talar": "قرآن تالار",
        "shirdarkola": "شیردارکلا",
        "sample": "نمونه اولیه (sample)",
    }
    heading(f"۹.{list(raw_pricing.keys()).index(village_key) + 1} روستا: {village_names.get(village_key, village_key)} "
            f"({len(items)} اقامتگاه)", 2)
    for item in items:
        d = item["data"]
        rid = item.get("room_id")
        pics = d.get("pictures") or []
        photo_descs = [
            f"{p.get('description', 'بدون توضیح')} ({p.get('created_at', '')})"
            for p in pics if p.get("description")
        ]
        rdate = room_dates.get(str(rid), {})
        heading(f"اقامتگاه {rid} — {d.get('title', '')}", 3)
        pairs = [
            ("شناسه", rid),
            ("وضعیت", d.get("status")),
            ("تعداد واحد", d.get("units_count")),
            ("آدرس", d.get("url")),
            ("آدرس کوتاه", d.get("short_url")),
            ("عنوان", d.get("title")),
            ("توضیحات", d.get("description")),
            ("عرض جغرافیایی", d.get("geo", {}).get("lat") if isinstance(d.get("geo"), dict) else None),
            ("طول جغرافیایی", d.get("geo", {}).get("lng") if isinstance(d.get("geo"), dict) else None),
            ("استان", d.get("province", {}).get("name") if isinstance(d.get("province"), dict) else None),
            ("شهر", d.get("city", {}).get("name") if isinstance(d.get("city"), dict) else None),
            ("نوع واگذاری", d.get("allocation")),
            ("نوع اقامتگاه", compact_list(d.get("types"))),
            ("منطقه", compact_list(d.get("regions"))),
            ("زیربنا (متر)", d.get("floor_area")),
            ("زمین (متر)", d.get("land_area")),
            ("تعداد طبقات", d.get("floors_count")),
            ("تعداد اتاق‌خواب", d.get("bedrooms")),
            ("چیدمان خواب", compact_list(d.get("sleep_arrange"))),
            ("شرح خواب", d.get("sleep_description")),
            ("ظرفیت عادی (نفر)", d.get("guest_number")),
            ("حداکثر ظرفیت (نفر)", d.get("max_guest_number")),
            ("حداقل اقامت (شب)", d.get("stays_min")),
            ("حداکثر اقامت (شب)", d.get("stays_max")),
            ("ساعت ورود", d.get("entrance_time_min")),
            ("ساعت ورود (تا)", d.get("entrance_time_max")),
            ("ساعت خروج", d.get("leaving_time")),
            ("امکانات", compact_list([f.get("name") for f in (d.get("features") or [])])),
            ("شرح امکانات", compact_dict({f.get("name"): f.get("description") for f in (d.get("features") or []) if f.get("description")})),
            ("ویژگی‌ها (برچسب‌ها)", compact_dict({p.get("name"): p.get("description") for p in (d.get("properties") or []) if p.get("name")})),
            ("امکانات اضافه", d.get("additional_feature")),
            ("ایمنی اضافه", d.get("additional_safety")),
            ("قوانین", compact_list(d.get("rules"))),
            ("قوانین اضافه", d.get("additional_rule")),
            ("سیاست لغو", d.get("cancellation_policy")),
            ("قیمت پایه هر شب (تومان)", d.get("min_price")),
            ("قیمت هر نفر اضافه (تومان)", d.get("extra_price")),
            ("نشان نظافت", d.get("is_clean")),
            ("اقامتگاه جدید", d.get("is_new")),
            ("رزرو فوری", d.get("is_instant")),
            ("اقامتگاه ممتاز", d.get("is_plus")),
            ("رزرو موفق", d.get("success_books")),
            ("تخفیف‌ها", compact_list(d.get("discounts"))),
            ("تخفیف فعلی", d.get("current_discount")),
            ("درصد تخفیف فعلی", d.get("current_discount_percent")),
            ("تعداد نظرات", d.get("ratings", {}).get("count") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز کلی", d.get("ratings", {}).get("total") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز دقت", d.get("ratings", {}).get("accuracy") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز ارتباط", d.get("ratings", {}).get("communication") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز نظافت", d.get("ratings", {}).get("cleanliness") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز موقعیت", d.get("ratings", {}).get("location") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز ورود", d.get("ratings", {}).get("checkin") if isinstance(d.get("ratings"), dict) else None),
            ("امتیاز ارزش", d.get("ratings", {}).get("value") if isinstance(d.get("ratings"), dict) else None),
            ("اخیراً فعال", d.get("ratings", {}).get("is_recently") if isinstance(d.get("ratings"), dict) else None),
            ("میزبان نامطلوب", d.get("ratings", {}).get("is_bad_host") if isinstance(d.get("ratings"), dict) else None),
            ("شناسه میزبان", d.get("host", {}).get("id") if isinstance(d.get("host"), dict) else None),
            ("نام میزبان", d.get("host", {}).get("name") if isinstance(d.get("host"), dict) else None),
            ("جنسیت میزبان", fa_gender(d.get("host", {}).get("gender")) if isinstance(d.get("host"), dict) else None),
            ("توضیح میزبان", d.get("host", {}).get("description") if isinstance(d.get("host"), dict) else None),
            ("عضویت میزبان", d.get("host", {}).get("created_at") if isinstance(d.get("host"), dict) else None),
            ("زمان پاسخ میزبان", d.get("host", {}).get("response_time") if isinstance(d.get("host"), dict) else None),
            ("نرخ پذیرش میزبان", d.get("host", {}).get("accept_rate") if isinstance(d.get("host"), dict) else None),
            ("امتیاز ارتباط میزبان", d.get("host", {}).get("host_communication_rate") if isinstance(d.get("host"), dict) else None),
            ("تعداد عکس", len(pics)),
            ("توضیحات عکس‌ها", compact_list(photo_descs) if photo_descs else "ندارد"),
            ("تاریخ اولین عکس", rdate.get("first_photo")),
            ("تاریخ اولین نظر", rdate.get("first_review")),
            ("تاریخ تخمینی اقامتگاه", rdate.get("est_date")),
            ("متادیتا (عنوان)", d.get("meta", {}).get("title") if isinstance(d.get("meta"), dict) else None),
            ("متادیتا (توضیح)", d.get("meta", {}).get("description") if isinstance(d.get("meta"), dict) else None),
        ]
        kv_table(pairs, font_size=8.5, widths=(6.5, 10.5))
        para("", after=3)

doc.add_page_break()

# ---------------------------------------------------------------- appendix 1
heading("پیوست ۱. توضیحات میزبان‌ها", 1)
para("توضیح متنی ۱۵۱ میزبان از data/hosts-babolkenar.json (میزبان‌های بدون توضیح ثبت نشده‌اند).")
for h in hosts:
    if h.get("description"):
        heading(f"میزبان {h.get('id')} — {h.get('name', '')}", 3)
        para(str(h.get("description")), size=9.5)

doc.add_page_break()

# ---------------------------------------------------------------- appendix 2
heading("پیوست ۲. اسنپ‌شات‌های تاریخی عرضه", 1)
para("عکس‌های نقطه‌ای از وضعیت عرضه در زمان‌های مختلف (از پوشه data/snapshots).")
for s in snapshots:
    heading(f"اسنپ‌شات {s.get('date', '')}", 3)
    kv_table(
        [
            ("تاریخ", s.get("date")),
            ("شمارنده‌ها", json.dumps(s.get("meta_counts", {}), ensure_ascii=False)),
            ("تعداد شناسه اقامتگاه‌ها", len(s.get("room_ids", []))),
        ],
        font_size=9,
    )

# ---------------------------------------------------------------- save
doc.save(OUT)
print(f"✅ Report saved: {OUT}")
print(f"   size: {OUT.stat().st_size / 1024:.0f} KB")