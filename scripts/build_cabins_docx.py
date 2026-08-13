# -*- coding: utf-8 -*-
"""Build Word docx: cabins in my price range, title cells = hyperlinks."""
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

MY_PRICE = 2750000
LO, HI = int(MY_PRICE * 0.75), int(MY_PRICE * 1.25)
MY_ID = 3297585
BASE = "https://www.jajiga.com/room/"

with open("data/pricing/pricing-dataset.json", encoding="utf-8") as f:
    cabins = json.load(f)

in_range = [c for c in cabins if LO <= c["min_price"] <= HI]
in_range.sort(key=lambda c: c["min_price"])


def village_from_title(title):
    m = re.search(r"بابلکنار\s*-\s*([^\s-]+)", title)
    return m.group(1) if m else "—"


def set_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "B Nazanin"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "B Nazanin")
    rFonts.set(qn("w:hAnsi"), "B Nazanin")
    rFonts.set(qn("w:cs"), "B Nazanin")
    bidi = OxmlElement("w:bidi")
    rPr.append(bidi)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_hyperlink(paragraph, url, text, size=11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "B Nazanin")
    rFonts.set(qn("w:hAnsi"), "B Nazanin")
    rFonts.set(qn("w:cs"), "B Nazanin")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    bidi = OxmlElement("w:bidi")
    rPr.append(bidi)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "B Nazanin"
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:ascii"), "B Nazanin")
rFonts.set(qn("w:hAnsi"), "B Nazanin")
rFonts.set(qn("w:cs"), "B Nazanin")

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"کلبه‌های رنج قیمتی من — {LO:,} تا {HI:,} تومان")
set_font(r, size=16, bold=True)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(f"({len(in_range)} کلبه در محدوده ±۲۵٪ قیمت کلبه سوئیسی سیدکلا — ۲,۷۵۰,۰۰۰ تومان)")
set_font(r2, size=11)

# Table
headers = ["ردیف", "عنوان کلبه (لینک)", "قیمت (تومان)", "مساحت (م²)", "خواب", "نفرات", "روستا"]
widths = [Cm(1.2), Cm(7.5), Cm(2.6), Cm(2.0), Cm(1.4), Cm(1.8), Cm(2.0)]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for i, (h, w) in enumerate(zip(headers, widths)):
    table.columns[i].width = w

hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].width = widths[i]
    para = hdr[i].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(h)
    set_font(run, size=11, bold=True)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "D9E2F3")
    hdr[i]._tc.get_or_add_tcPr().append(shd)

for idx, c in enumerate(in_range, 1):
    row = table.add_row().cells
    vals = [
        str(idx),
        c["title"],
        f"{c['min_price']:,}",
        str(c.get("floor_area") or "—"),
        str(c.get("bedrooms") or "—"),
        f"{c.get('guest_number') or '—'}-{c.get('max_guest_number') or '—'}",
        village_from_title(c["title"]),
    ]
    for i, v in enumerate(vals):
        row[i].width = widths[i]
        para = row[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 1:
            add_hyperlink(para, BASE + str(c["id"]), v)
        else:
            run = para.add_run(v)
            set_font(run, size=10)

out = r"C:\Users\Ma\Desktop\کلبه های رنج قیمتی من.docx"
doc.save(out)
print("Saved:", out)
print("Rows:", len(in_range))
